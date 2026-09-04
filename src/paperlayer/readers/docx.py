"""DOCX reader.

python-docx is used to open the package and hand over the XML; the structural
work is done here. Three things in particular are resolved by hand because the
library does not expose them in usable form:

* **Style inheritance.** ``w:basedOn`` chains are walked so that a run with no
  explicit size still reports the size it will actually render at, which the
  fallback heading detector needs.
* **Numbering.** ``numbering.xml`` is read to map a paragraph numbering
  reference to a real list format and nesting level. Without it every list is
  an undifferentiated paragraph.
* **Body order.** ``document.paragraphs`` and ``document.tables`` are separate
  sequences that lose their interleaving; the body element is walked directly
  so a table stays between the paragraphs that surround it.

DOCX carries explicit structure, so this reader emits blocks directly rather
than going through the geometric pipeline the PDF path needs.
"""

from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass, field
from typing import Any, BinaryIO

from .._text import collapse_ws, normalize_unicode, split_list_marker
from ..errors import MissingBackendError, ParseError
from ..options import ParseOptions
from ..pipeline.captions import caption_label
from ..pipeline.headings import is_heading_candidate, normalize_levels, rank_levels
from ..pipeline.tables import looks_like_header
from ..types import Block, Footnote, SourceInfo, StyleInfo, TableData

__all__ = ["read_docx"]

#: Word stores font sizes in half-points.
_HALF_POINT = 2.0
#: Used when a document specifies no size anywhere.
_DEFAULT_SIZE = 11.0

#: Heading style names across the localisations that turn up most often.
_HEADING_RE = re.compile(
    r"^(?:heading|überschrift|uberschrift|titre|t[ií]tulo|kop|rubrik|otsikko)"
    r"\s*(\d)\b",
    re.IGNORECASE,
)
_TITLE_STYLES = {"title", "titel", "titre", "título", "subtitle", "untertitel"}


@dataclass(slots=True)
class _StyleSpec:
    """Effective formatting of a named style, after inheritance."""

    name: str = ""
    size: float | None = None
    bold: bool | None = None
    outline: int | None = None
    based_on: str | None = None


@dataclass(slots=True)
class _Para:
    """One paragraph, flattened to what classification needs."""

    text: str
    size: float
    bold: bool
    style: str
    outline: int | None
    num_id: str | None = None
    num_level: int = 0
    num_format: str = "decimal"
    refs: list[str] = field(default_factory=list)


def read_docx(
    source: str | os.PathLike[str] | bytes | BinaryIO,
    options: ParseOptions,
) -> tuple[list[Block], SourceInfo, list[str]]:
    """Read a DOCX into blocks, its source metadata and any warnings."""
    try:
        import docx
        from docx.oxml.ns import qn
    except ImportError as exc:  # pragma: no cover - depends on install shape
        raise MissingBackendError("python-docx", "docx") from exc

    stream: Any = io.BytesIO(source) if isinstance(source, bytes) else source
    path = (
        os.fspath(source)
        if isinstance(source, (str, os.PathLike))
        else getattr(source, "name", None)
    )

    try:
        document = docx.Document(stream)
    except Exception as exc:
        raise ParseError(f"Could not open the DOCX: {exc}") from exc

    warnings: list[str] = []
    styles, default_size = _style_table(document, qn)
    numbering = _numbering_table(document, qn, warnings)
    notes, note_order = _note_table(document, qn, warnings)

    paragraphs: list[_Para] = []
    body_items: list[tuple[str, Any]] = []

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para = _read_paragraph(child, qn, styles, default_size, numbering, note_order)
            if para is not None:
                body_items.append(("p", para))
                paragraphs.append(para)
        elif child.tag == qn("w:tbl"):
            body_items.append(("tbl", child))

    blocks = _build_blocks(body_items, paragraphs, notes, qn, options)

    props = document.core_properties
    info = SourceInfo(
        format="docx",
        path=path,
        n_pages=None,
        title=(props.title or None),
        author=(props.author or None),
        producer=(props.last_modified_by or None),
    )
    return blocks, info, warnings


# --------------------------------------------------------------------------
# styles.xml
# --------------------------------------------------------------------------


def _style_table(document: Any, qn: Any) -> tuple[dict[str, _StyleSpec], float]:
    """Resolve every style to its effective size, weight and outline level."""
    raw: dict[str, _StyleSpec] = {}
    default_size = _DEFAULT_SIZE

    try:
        root = document.styles.element
    except Exception:  # pragma: no cover - malformed package
        return raw, default_size

    defaults = root.find(qn("w:docDefaults"))
    if defaults is not None:
        rpr = defaults.find(qn("w:rPrDefault"))
        if rpr is not None:
            size = _run_size(rpr.find(qn("w:rPr")), qn)
            if size is not None:
                default_size = size

    for element in root.findall(qn("w:style")):
        style_id = element.get(qn("w:styleId"))
        if not style_id:
            continue
        name_el = element.find(qn("w:name"))
        based_el = element.find(qn("w:basedOn"))
        ppr = element.find(qn("w:pPr"))
        outline = None
        if ppr is not None:
            outline_el = ppr.find(qn("w:outlineLvl"))
            if outline_el is not None:
                outline = _int_val(outline_el.get(qn("w:val")))
        raw[style_id] = _StyleSpec(
            name=(name_el.get(qn("w:val")) if name_el is not None else style_id) or style_id,
            size=_run_size(element.find(qn("w:rPr")), qn),
            bold=_run_bold(element.find(qn("w:rPr")), qn),
            outline=outline,
            based_on=(based_el.get(qn("w:val")) if based_el is not None else None),
        )

    return {sid: _resolve(sid, raw) for sid in raw}, default_size


def _resolve(style_id: str, raw: dict[str, _StyleSpec]) -> _StyleSpec:
    """Flatten a ``w:basedOn`` chain, guarding against cycles."""
    spec = raw[style_id]
    size, bold, outline = spec.size, spec.bold, spec.outline
    seen = {style_id}
    parent_id = spec.based_on
    while parent_id and parent_id in raw and parent_id not in seen:
        seen.add(parent_id)
        parent = raw[parent_id]
        if size is None:
            size = parent.size
        if bold is None:
            bold = parent.bold
        if outline is None:
            outline = parent.outline
        parent_id = parent.based_on
    return _StyleSpec(name=spec.name, size=size, bold=bold, outline=outline)


def _run_size(rpr: Any, qn: Any) -> float | None:
    if rpr is None:
        return None
    element = rpr.find(qn("w:sz"))
    if element is None:
        return None
    value = _int_val(element.get(qn("w:val")))
    return value / _HALF_POINT if value else None


def _run_bold(rpr: Any, qn: Any) -> bool | None:
    if rpr is None:
        return None
    element = rpr.find(qn("w:b"))
    if element is None:
        return None
    value = element.get(qn("w:val"))
    return value not in ("0", "false", "off")


def _int_val(value: str | None) -> int | None:
    try:
        return int(value) if value is not None else None
    except ValueError:
        return None


# --------------------------------------------------------------------------
# numbering.xml
# --------------------------------------------------------------------------


def _numbering_table(document: Any, qn: Any, warnings: list[str]) -> dict[tuple[str, int], str]:
    """Map ``(numId, ilvl)`` to a list format such as ``bullet`` or ``decimal``."""
    try:
        part = document.part.numbering_part
        root = part.element
    except Exception:
        return {}

    abstract: dict[str, dict[int, str]] = {}
    for node in root.findall(qn("w:abstractNum")):
        abstract_id = node.get(qn("w:abstractNumId"))
        if abstract_id is None:
            continue
        levels: dict[int, str] = {}
        for lvl in node.findall(qn("w:lvl")):
            index = _int_val(lvl.get(qn("w:ilvl"))) or 0
            fmt_el = lvl.find(qn("w:numFmt"))
            levels[index] = (
                fmt_el.get(qn("w:val")) if fmt_el is not None else "decimal"
            ) or "decimal"
        abstract[abstract_id] = levels

    mapping: dict[tuple[str, int], str] = {}
    for node in root.findall(qn("w:num")):
        num_id = node.get(qn("w:numId"))
        ref = node.find(qn("w:abstractNumId"))
        if num_id is None or ref is None:
            continue
        levels = abstract.get(ref.get(qn("w:val")) or "", {})
        for index, fmt in levels.items():
            mapping[(num_id, index)] = fmt
    if not mapping and abstract:
        warnings.append("numbering.xml defined no usable list levels")
    return mapping


# --------------------------------------------------------------------------
# footnotes.xml / endnotes.xml
# --------------------------------------------------------------------------


def _note_table(
    document: Any, qn: Any, warnings: list[str]
) -> tuple[dict[str, str], dict[str, str]]:
    """Footnote and endnote text by id, plus a mutable display-number map.

    Word keeps notes in a separate part with ids that are not the numbers the
    reader sees; the display map is filled in as references are encountered, so
    markers come out 1, 2, 3 in reading order.
    """
    notes: dict[str, str] = {}
    for relationship in document.part.rels.values():
        reltype = str(relationship.reltype)
        if not (reltype.endswith("/footnotes") or reltype.endswith("/endnotes")):
            continue
        prefix = "e" if reltype.endswith("/endnotes") else ""
        try:
            root = _parse_part(relationship.target_part)
        except Exception as exc:
            warnings.append(f"could not read notes part: {exc}")
            continue
        if root is None:
            continue
        tag = qn("w:endnote") if prefix else qn("w:footnote")
        for node in root.findall(tag):
            note_type = node.get(qn("w:type"))
            if note_type in ("separator", "continuationSeparator", "continuationNotice"):
                continue
            note_id = node.get(qn("w:id"))
            if note_id is None:
                continue
            text = collapse_ws(" ".join(t.text or "" for t in node.iter(qn("w:t"))))
            # Word writes the reference mark into the note itself; drop it so
            # the marker is not duplicated in the rendered output.
            text = re.sub(r"^[\d*†‡§¶]{1,3}[.)]?\s*", "", text).strip()
            if text:
                notes[prefix + note_id] = text
    return notes, {}


def _parse_part(part: Any) -> Any:
    element = getattr(part, "element", None)
    if element is not None:
        return element
    from lxml import etree  # python-docx already depends on lxml

    return etree.fromstring(part.blob)


# --------------------------------------------------------------------------
# Paragraphs
# --------------------------------------------------------------------------


def _read_paragraph(
    element: Any,
    qn: Any,
    styles: dict[str, _StyleSpec],
    default_size: float,
    numbering: dict[tuple[str, int], str],
    note_order: dict[str, str],
) -> _Para | None:
    ppr = element.find(qn("w:pPr"))
    style_id = None
    if ppr is not None:
        style_el = ppr.find(qn("w:pStyle"))
        if style_el is not None:
            style_id = style_el.get(qn("w:val"))

    spec = styles.get(style_id or "", _StyleSpec())
    style_size = spec.size if spec.size is not None else default_size
    style_bold = bool(spec.bold)

    text, refs = _paragraph_text(element, qn, note_order)
    text = collapse_ws(normalize_unicode(text))
    if not text:
        return None

    size, bold = _paragraph_typography(element, qn, style_size, style_bold)

    num_id: str | None = None
    num_level = 0
    num_format = "decimal"
    if ppr is not None:
        numpr = ppr.find(qn("w:numPr"))
        if numpr is not None:
            id_el = numpr.find(qn("w:numId"))
            lvl_el = numpr.find(qn("w:ilvl"))
            num_id = id_el.get(qn("w:val")) if id_el is not None else None
            num_level = _int_val(lvl_el.get(qn("w:val")) if lvl_el is not None else None) or 0
            if num_id is not None:
                num_format = numbering.get((num_id, num_level), "decimal")

    outline = spec.outline
    if ppr is not None:
        outline_el = ppr.find(qn("w:outlineLvl"))
        if outline_el is not None:
            outline = _int_val(outline_el.get(qn("w:val")))

    return _Para(
        text=text,
        size=size,
        bold=bold,
        style=spec.name or (style_id or ""),
        outline=outline,
        num_id=num_id,
        num_level=num_level,
        num_format=num_format,
        refs=refs,
    )


def _paragraph_text(element: Any, qn: Any, note_order: dict[str, str]) -> tuple[str, list[str]]:
    """Flatten a paragraph to text, turning note references into ``[^n]``.

    Walking the element tree in document order is what keeps the reference in
    the right place inside the sentence rather than appended at the end.
    """
    t_tag = qn("w:t")
    tab_tag = qn("w:tab")
    br_tag = qn("w:br")
    fn_tag = qn("w:footnoteReference")
    en_tag = qn("w:endnoteReference")

    parts: list[str] = []
    refs: list[str] = []
    for node in element.iter():
        tag = node.tag
        if tag == t_tag:
            parts.append(node.text or "")
        elif tag in (tab_tag, br_tag):
            parts.append(" ")
        elif tag in (fn_tag, en_tag):
            note_id = node.get(qn("w:id"))
            if note_id is None:
                continue
            key = ("e" if tag == en_tag else "") + note_id
            marker = note_order.get(key)
            if marker is None:
                marker = str(len(note_order) + 1)
                note_order[key] = marker
            parts.append(f"[^{marker}]")
            refs.append(key)
    return "".join(parts), refs


def _paragraph_typography(
    element: Any, qn: Any, style_size: float, style_bold: bool
) -> tuple[float, bool]:
    """Character-weighted size and weight, with run overrides beating the style."""
    sizes: dict[float, int] = {}
    bold_chars = 0
    total = 0
    for run in element.iter(qn("w:r")):
        text = "".join(t.text or "" for t in run.iter(qn("w:t")))
        length = len(text)
        if not length:
            continue
        rpr = run.find(qn("w:rPr"))
        size = _run_size(rpr, qn)
        size = style_size if size is None else size
        bold = _run_bold(rpr, qn)
        bold = style_bold if bold is None else bold
        sizes[round(size, 1)] = sizes.get(round(size, 1), 0) + length
        total += length
        if bold:
            bold_chars += length

    if not total:
        return style_size, style_bold
    size = max(sizes.items(), key=lambda kv: (kv[1], kv[0]))[0]
    return size, (bold_chars / total) >= 0.6


# --------------------------------------------------------------------------
# Classification
# --------------------------------------------------------------------------


def _build_blocks(
    body_items: list[tuple[str, Any]],
    paragraphs: list[_Para],
    notes: dict[str, str],
    qn: Any,
    options: ParseOptions,
) -> list[Block]:
    body = _body_size(paragraphs)
    fallback = _fallback_levels(paragraphs, body, options)

    blocks: list[Block] = []
    pending: list[_Para] = []

    def flush_list() -> None:
        if pending:
            blocks.append(_list_block(pending, body))
            pending.clear()

    for kind, item in body_items:
        if kind == "tbl":
            flush_list()
            block = _table_block(item, qn)
            if block is not None:
                blocks.append(block)
            continue

        para: _Para = item
        level = _heading_level(para, options)
        if level is None and fallback is not None:
            level = fallback.get(id(para))

        if level is not None:
            flush_list()
            blocks.append(
                Block(
                    type="heading",
                    text=para.text,
                    level=min(level, options.max_heading_level),
                    style=_style_info(para, body),
                    meta={"source_style": para.style} if para.style else {},
                )
            )
            continue

        if options.detect_lists and _is_list_item(para):
            if pending and pending[-1].num_id != para.num_id:
                # A different numbering definition is a different list,
                # even with no paragraph between them.
                flush_list()
            pending.append(para)
            continue
        flush_list()

        block_type = "paragraph"
        meta: dict[str, Any] = {}
        if options.detect_captions:
            if "caption" in para.style.lower():
                block_type = "caption"
            label = caption_label(para.text)
            if label is not None:
                kind_name, name = label
                block_type = "caption"
                meta = {"caption_kind": kind_name, "label": name}
        blocks.append(
            Block(
                type=block_type,  # type: ignore[arg-type]
                text=para.text,
                style=_style_info(para, body),
                meta=meta,
            )
        )

    flush_list()

    _renumber(blocks, options)
    if options.detect_footnotes:
        _attach_notes(blocks, notes)
    return blocks


def _body_size(paragraphs: list[_Para]) -> float:
    weights: dict[float, int] = {}
    for para in paragraphs:
        weights[round(para.size, 1)] = weights.get(round(para.size, 1), 0) + len(para.text)
    if not weights:
        return _DEFAULT_SIZE
    return max(weights.items(), key=lambda kv: (kv[1], -kv[0]))[0]


def _style_info(para: _Para, body: float) -> StyleInfo:
    return StyleInfo(
        size=para.size,
        size_ratio=round(para.size / body, 3) if body else 1.0,
        bold_ratio=1.0 if para.bold else 0.0,
        italic_ratio=0.0,
        font=None,
    )


def _heading_level(para: _Para, options: ParseOptions) -> int | None:
    """Level from the named style or outline level, the authoritative signals."""
    if not options.detect_headings:
        return None
    name = para.style.strip().lower()
    match = _HEADING_RE.match(name)
    if match:
        return max(1, min(int(match.group(1)), options.max_heading_level))
    if name in _TITLE_STYLES:
        return 1 if name not in ("subtitle", "untertitel") else 2
    if para.outline is not None and 0 <= para.outline <= 8:
        return para.outline + 1
    return None


def _fallback_levels(
    paragraphs: list[_Para], body: float, options: ParseOptions
) -> dict[int, int] | None:
    """Font-metric heading detection, for documents that use no heading styles.

    Plenty of real-world DOCX files are typed by hand with bold, larger text
    instead of styles; without this they would come out as one flat wall of
    paragraphs.
    """
    if not options.detect_headings:
        return None
    if any(_heading_level(p, options) is not None for p in paragraphs):
        return None

    bold_share = sum(1 for p in paragraphs if p.bold) / len(paragraphs) if paragraphs else 0.0
    saturated = bold_share > 0.55

    candidates = [
        p
        for p in paragraphs
        if not _is_list_item(p)
        and is_heading_candidate(p.text, p.size, p.bold, body, saturated, options)
    ]
    if not candidates:
        return None

    levels = rank_levels({(round(p.size, 1), p.bold) for p in candidates}, options)
    return {id(p): levels[(round(p.size, 1), p.bold)] for p in candidates}


def _is_list_item(para: _Para) -> bool:
    if para.num_id is not None:
        return True
    # A "List Paragraph" style with a literal bullet typed in is still a list.
    return "list" in para.style.lower() and split_list_marker(para.text) is not None


def _list_block(items: list[_Para], body: float) -> Block:
    payload: list[dict[str, Any]] = []
    for para in items:
        text = para.text
        marker = "-"
        ordered = para.num_format not in ("bullet", "none")
        if para.num_id is None:
            split = split_list_marker(text)
            if split is not None:
                marker, text, ordered = split
        payload.append(
            {
                "text": text,
                "level": para.num_level,
                "marker": marker,
                "ordered": ordered,
            }
        )

    text = "\n".join(
        "  " * int(entry["level"]) + ("1. " if entry["ordered"] else "- ") + str(entry["text"])
        for entry in payload
    )
    return Block(
        type="list",
        text=text,
        level=max(int(entry["level"]) for entry in payload) + 1,
        style=_style_info(items[0], body),
        meta={"items": payload},
    )


def _table_block(element: Any, qn: Any) -> Block | None:
    """Read a table straight from the XML, expanding merges to a plain grid.

    The row/cell API of python-docx repeats the same cell object across a
    merged span, and telling those repeats apart by object identity is not
    safe: lxml creates proxy objects on demand, so two different cells can
    share an ``id()`` once the first proxy has been collected. Reading
    ``w:gridSpan`` and ``w:vMerge`` directly is both correct and simpler.
    """
    rows: list[list[str]] = []
    bold_ratios: list[float] = []

    for tr in element.findall(qn("w:tr")):
        values: list[str] = []
        bold_chars = 0
        total = 0
        for tc in tr.findall(qn("w:tc")):
            span, continues = _cell_span(tc, qn)
            if continues:
                # A vertically merged cell repeats its content downwards; the
                # text belongs to the row that started the merge.
                values.extend([""] * span)
                continue
            text = collapse_ws(
                normalize_unicode(" ".join(t.text or "" for t in tc.iter(qn("w:t"))))
            )
            values.append(text)
            values.extend([""] * (span - 1))
            for run in tc.iter(qn("w:r")):
                run_text = "".join(t.text or "" for t in run.iter(qn("w:t")))
                total += len(run_text)
                if _run_bold(run.find(qn("w:rPr")), qn):
                    bold_chars += len(run_text)
        if values:
            rows.append(values)
            bold_ratios.append(bold_chars / total if total else 0.0)

    if not rows or not any(cell for row in rows for cell in row):
        return None

    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]

    has_header = _explicit_header(element, qn) or looks_like_header(rows, bold_ratios)
    header = rows[0] if has_header else None
    data = rows[1:] if has_header else rows

    return Block(
        type="table",
        text="\n".join("\t".join(row) for row in rows),
        table=TableData(rows=data, header=header, ruled=True),
        meta={"ruled": True, "n_rows": len(rows), "n_cols": width},
    )


def _cell_span(tc: Any, qn: Any) -> tuple[int, bool]:
    """``(grid_span, is_vertical_continuation)`` for one table cell."""
    tcpr = tc.find(qn("w:tcPr"))
    if tcpr is None:
        return 1, False
    span = 1
    grid = tcpr.find(qn("w:gridSpan"))
    if grid is not None:
        span = _int_val(grid.get(qn("w:val"))) or 1
    merge = tcpr.find(qn("w:vMerge"))
    continues = merge is not None and (merge.get(qn("w:val")) or "continue") != "restart"
    return max(span, 1), continues


def _explicit_header(element: Any, qn: Any) -> bool:
    """Whether the first row is marked as a repeating header row in Word."""
    first = element.find(qn("w:tr"))
    if first is None:
        return False
    trpr = first.find(qn("w:trPr"))
    if trpr is None:
        return False
    element = trpr.find(qn("w:tblHeader"))
    return element is not None and element.get(qn("w:val")) not in ("0", "false", "off")


def _renumber(blocks: list[Block], options: ParseOptions) -> None:
    headings = [b for b in blocks if b.type == "heading"]
    if headings:
        levels = normalize_levels([b.level or 1 for b in headings], options)
        for block, level in zip(headings, levels, strict=True):
            block.level = level
    for index, block in enumerate(blocks):
        block.order = index


def _attach_notes(blocks: list[Block], notes: dict[str, str]) -> None:
    """Attach each note to the block whose text carries its ``[^n]`` marker."""
    if not notes:
        return
    # note_order mapped part-ids to display markers in reading order; rebuild
    # that association by walking the markers as they appear in the blocks.
    by_marker: dict[str, str] = {}
    ordered_ids = list(notes.keys())
    for index, note_id in enumerate(sorted(ordered_ids, key=_note_sort_key), start=1):
        by_marker[str(index)] = notes[note_id]

    for block in blocks:
        for marker in re.findall(r"\[\^([^\]]{1,3})\]", block.text):
            text = by_marker.get(marker)
            if text is not None:
                block.footnotes.append(Footnote(marker=marker, text=text))


def _note_sort_key(note_id: str) -> tuple[int, int]:
    endnote = note_id.startswith("e")
    digits = note_id[1:] if endnote else note_id
    try:
        value = int(digits)
    except ValueError:
        value = 0
    return (1 if endnote else 0, value)

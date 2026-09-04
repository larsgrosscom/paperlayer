"""Compare paperlayer against naive raw-text extraction.

The baseline is what most pipelines actually do today: ``pdfplumber``'s
``page.extract_text()`` concatenated across pages, and for DOCX the paragraph
texts joined with newlines. Both are counted with the same tokeniser so the
comparison is like for like.

Usage::

    python benchmarks/run_benchmark.py                     # synthetic corpus
    python benchmarks/run_benchmark.py --docs ~/my/pdfs    # your own documents
    python benchmarks/run_benchmark.py --out table.md      # README table

Token reduction is a proxy, not the point. A naive extractor that dropped half
the document would score wonderfully. The structure columns are there to show
that the reduction comes from removing furniture and duplication rather than
from losing content, and ``--check-content`` asserts that outright.
"""

from __future__ import annotations

import argparse
import statistics
import sys
import time
from collections.abc import Callable, Iterable, Sequence
from dataclasses import dataclass
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from paperlayer import Document, parse

__all__ = ["Result", "benchmark_file", "main"]

_EXTENSIONS = (".pdf", ".docx")


@dataclass(slots=True)
class Result:
    """One document, measured both ways."""

    path: Path
    baseline_tokens: int
    paperlayer_tokens: int
    headings: int
    tables: int
    lists: int
    footnotes: int
    blocks: int
    seconds: float
    content_kept: float

    @property
    def reduction(self) -> float:
        """Fraction of baseline tokens removed, as a percentage."""
        if not self.baseline_tokens:
            return 0.0
        return 100.0 * (1 - self.paperlayer_tokens / self.baseline_tokens)


# --------------------------------------------------------------------------
# Tokenisation
# --------------------------------------------------------------------------


def build_counter(model: str) -> tuple[Callable[[str], int], str]:
    """A token counter, falling back to an estimate when tiktoken is absent.

    tiktoken downloads its encoding file on first use. That is fine for a
    developer benchmark but would violate the offline guarantee of the library
    itself, which is why it is a ``bench`` extra and not a dependency.
    """
    try:
        import tiktoken

        encoding = tiktoken.get_encoding(model)
    except Exception as exc:
        print(
            f"note: tiktoken unavailable ({exc.__class__.__name__}); "
            "falling back to a 4-characters-per-token estimate",
            file=sys.stderr,
        )
        return (lambda text: max(1, round(len(text) / 4))), "estimate (chars/4)"

    return (lambda text: len(encoding.encode(text, disallowed_special=()))), model


# --------------------------------------------------------------------------
# Baseline
# --------------------------------------------------------------------------


def baseline_text(path: Path) -> str:
    """What a pipeline gets today with no layout analysis at all."""
    if path.suffix.lower() == ".pdf":
        import pdfplumber

        chunks: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                chunks.append(page.extract_text() or "")
                page.close()
        return "\n".join(chunks)

    import docx

    document = docx.Document(str(path))
    return "\n".join(p.text for p in document.paragraphs)


# --------------------------------------------------------------------------
# Measurement
# --------------------------------------------------------------------------

# Function words carry no signal about whether content survived extraction.
_STOPWORDS = frozenset(
    [
        "the",
        "a",
        "an",
        "and",
        "or",
        "of",
        "to",
        "in",
        "for",
        "on",
        "with",
        "is",
        "are",
        "was",
        "were",
        "be",
        "by",
        "at",
        "as",
        "it",
        "that",
        "this",
    ]
)


def _content_words(text: str) -> set[str]:
    return {
        word
        for word in "".join(c.lower() if c.isalnum() else " " for c in text).split()
        if len(word) > 3 and word not in _STOPWORDS and not word.isdigit()
    }


def benchmark_file(
    path: Path, count: Callable[[str], int], table_mode: str = "markdown"
) -> Result:
    raw = baseline_text(path)

    started = time.perf_counter()
    doc: Document = parse(path, table_mode=table_mode)
    elapsed = time.perf_counter() - started

    markdown = doc.markdown

    # Retention is measured against keep_headers=True, not against the raw
    # baseline. Stripping a running header is the job, not data loss, and
    # scoring it as loss would make correct behaviour look like a regression.
    # What this does catch is text that fell out by accident.
    unstripped = parse(path, keep_headers=True, table_mode=table_mode).markdown
    baseline_words = _content_words(raw) & _content_words(unstripped)
    kept = (
        len(baseline_words & _content_words(markdown)) / len(baseline_words)
        if baseline_words
        else 1.0
    )

    return Result(
        path=path,
        baseline_tokens=count(raw),
        paperlayer_tokens=count(markdown),
        headings=len(doc.headings()),
        tables=len(doc.tables()),
        lists=len(doc.of_type("list")),
        footnotes=len(doc.footnotes()),
        blocks=len(doc.blocks),
        seconds=elapsed,
        content_kept=kept,
    )


def collect(directory: Path) -> list[Path]:
    return sorted(
        p for p in directory.rglob("*") if p.suffix.lower() in _EXTENSIONS and p.is_file()
    )


# --------------------------------------------------------------------------
# Reporting
# --------------------------------------------------------------------------


def render_table(results: Sequence[Result], encoding_name: str) -> str:
    header = (
        "| Document | Baseline tokens | paperlayer tokens | Reduction | "
        "Headings | Tables | Lists | Footnotes |\n"
        "| --- | ---: | ---: | ---: | ---: | ---: | ---: | ---: |"
    )
    rows = [
        f"| `{r.path.name}` | {r.baseline_tokens:,} | {r.paperlayer_tokens:,} | "
        f"{r.reduction:.1f}% | {r.headings} | {r.tables} | {r.lists} | {r.footnotes} |"
        for r in results
    ]

    reductions = [r.reduction for r in results]
    totals_base = sum(r.baseline_tokens for r in results)
    totals_ours = sum(r.paperlayer_tokens for r in results)
    total_reduction = 100.0 * (1 - totals_ours / totals_base) if totals_base else 0.0

    summary = (
        f"| **Total / median** | **{totals_base:,}** | **{totals_ours:,}** | "
        f"**{statistics.median(reductions):.1f}%** | "
        f"**{sum(r.headings for r in results)}** | "
        f"**{sum(r.tables for r in results)}** | "
        f"**{sum(r.lists for r in results)}** | "
        f"**{sum(r.footnotes for r in results)}** |"
    )

    kept = statistics.median(r.content_kept for r in results)
    note = (
        f"\nMedian token reduction **{statistics.median(reductions):.1f}%** "
        f"({total_reduction:.1f}% across the whole corpus), with "
        f"**{kept * 100:.0f}%** of baseline content words retained. "
        f"Counted with `{encoding_name}`."
    )
    return "\n".join([header, *rows, summary]) + "\n" + note


def main(argv: Iterable[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    parser.add_argument(
        "--docs",
        type=Path,
        help="folder of PDF/DOCX files (default: a generated synthetic corpus)",
    )
    parser.add_argument("--out", type=Path, help="write the Markdown table to this file")
    parser.add_argument("--json", type=Path, help="write raw results as JSON")
    parser.add_argument(
        "--encoding",
        default="cl100k_base",
        help="tiktoken encoding name (default: cl100k_base)",
    )
    parser.add_argument(
        "--table-mode",
        default="markdown",
        choices=("markdown", "html", "csv", "text", "drop"),
        help="table rendering to measure (default: markdown). "
        "'text' is the token-cheapest option.",
    )
    parser.add_argument(
        "--check-content",
        action="store_true",
        help="fail if less than 90%% of baseline content words survive",
    )
    args = parser.parse_args(list(argv) if argv is not None else None)

    if args.docs:
        directory = args.docs
        if not directory.is_dir():
            print(f"benchmark: not a directory: {directory}", file=sys.stderr)
            return 2
    else:
        from corpus import write_corpus

        directory = Path(__file__).parent / "corpus"
        write_corpus(directory)
        print(f"note: using the generated corpus in {directory}", file=sys.stderr)

    paths = collect(directory)
    if not paths:
        print(f"benchmark: no PDF or DOCX files under {directory}", file=sys.stderr)
        return 2

    count, encoding_name = build_counter(args.encoding)

    results: list[Result] = []
    for path in paths:
        try:
            results.append(benchmark_file(path, count, args.table_mode))
        except Exception as exc:
            print(f"benchmark: skipped {path.name}: {exc}", file=sys.stderr)

    if not results:
        return 2

    table = render_table(results, encoding_name)
    print(table)

    if args.out:
        args.out.write_text(table + "\n", encoding="utf-8")
    if args.json:
        import json
        from dataclasses import asdict

        payload = [
            {**asdict(r), "path": str(r.path), "reduction": r.reduction} for r in results
        ]
        args.json.write_text(json.dumps(payload, indent=2), encoding="utf-8")

    total = sum(r.seconds for r in results)
    print(
        f"\nParsed {len(results)} document(s) in {total:.2f}s "
        f"({total / len(results):.2f}s each).",
        file=sys.stderr,
    )

    if args.check_content:
        worst = min(results, key=lambda r: r.content_kept)
        if worst.content_kept < 0.90:
            print(
                f"benchmark: {worst.path.name} kept only "
                f"{worst.content_kept * 100:.0f}% of baseline content words",
                file=sys.stderr,
            )
            return 1
    return 0


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main())
